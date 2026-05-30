"""
Relatório job handler — Fase 5.

Gera DOCX (python-docx) com estrutura do relatório padrão ScanSOLO:
  Capa → Sumário → Apresentação → Objetivo → Metodologia →
  Levantamento → Resultados (por alvo) → Conclusão

TODO (pós-validação com Marcos):
  - Converter DOCX → PDF automaticamente (LibreOffice headless ou DOCX2PDF)
  - Upload para Dropbox (via Dropbox client na Fase de produção)
  - Substituir boilerplate provisório dos textos padrão ScanSOLO
"""

from __future__ import annotations

import io
from datetime import date
from typing import TYPE_CHECKING, Any

import structlog
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

if TYPE_CHECKING:
    from clients.supabase_client import SupabaseClient

log = structlog.get_logger()

STORAGE_BUCKET = "gpr-tabelas"
SUPABASE_URL = __import__("os").environ.get("NEXT_PUBLIC_SUPABASE_URL", "")

# ── Mapeamento de tipo interno → rótulo para o relatório ───────────────────
_TIPO_LABEL: dict[str, str] = {
    "tubulacao_agua": "Tubulação de Água",
    "tubulacao_gas": "Tubulação de Gás",
    "cabo_eletrico": "Cabo Elétrico",
    "cabo_telecom": "Cabo de Telecomunicações",
    "vazio": "Vazio / Cavidade",
    "raiz": "Raiz",
    "rocha": "Rocha / Matacão",
    "desconhecido": "Interferência Não Identificada",
}

# ── Textos boilerplate (TODO: substituir com versão final aprovada pela ScanSOLO)
_BOILERPLATE_GPR = (
    "O Georadar, também conhecido como Ground Penetrating Radar (GPR), é um método geofísico "
    "de alta resolução que utiliza ondas eletromagnéticas de radiofrequência para investigar o "
    "interior do terreno de forma não destrutiva. As ondas são emitidas por uma antena "
    "transmissora e penetram no solo, sendo refletidas nas interfaces entre materiais com "
    "diferentes propriedades dielétricas. O tempo de retorno das reflexões é registrado e "
    "processado para gerar perfis bidimensionais (radargramas) que permitem identificar e "
    "mapear estruturas e interferências subterrâneas com precisão centimétrica em profundidade."
)

_BOILERPLATE_PIPE_LOCATOR = (
    "O Pipe Locator é um equipamento eletromagnético utilizado para localizar e mapear redes "
    "de tubulações e cabos enterrados condutivos. O método baseia-se na detecção do campo "
    "eletromagnético gerado pela passagem de corrente elétrica (ativa ou induzida) pelo "
    "condutor metálico. O equipamento é composto por um transmissor, que injeta ou induz "
    "sinal de radiofrequência no condutor, e um receptor portátil, que detecta o campo "
    "eletromagnético resultante. A técnica é complementar ao GPR, sendo especialmente eficaz "
    "para localização de tubulações metálicas, cabos elétricos e de telecomunicações."
)

_BOILERPLATE_DISCLAIMER = (
    "Os resultados apresentados a seguir foram obtidos a partir do processamento e "
    "interpretação dos dados coletados em campo. As profundidades e diâmetros indicados "
    "são estimativas baseadas nos parâmetros de velocidade de propagação adotados, podendo "
    "apresentar variações em função das condições locais do solo. Recomenda-se a confirmação "
    "dos resultados por meio de sondagem ou escavação cuidadosa antes de qualquer intervenção."
)

_BOILERPLATE_CONCLUSAO = (
    "Com base nos resultados obtidos pelo levantamento geofísico com Georadar (GPR) e "
    "Pipe Locator, foi possível identificar e mapear as interferências subterrâneas presentes "
    "na área investigada. As interferências identificadas foram classificadas por tipo, "
    "profundidade estimada e diâmetro aparente, conforme detalhado na seção de Resultados "
    "deste relatório.\n\n"
    "Recomenda-se que qualquer escavação ou intervenção na área seja precedida de inspeção "
    "visual e localização precisa das interferências identificadas, utilizando técnicas não "
    "destrutivas adicionais quando necessário. A ScanSOLO disponibiliza-se para "
    "esclarecimentos adicionais sobre os resultados apresentados neste relatório."
)

_RODAPE_FIXO = "ScanSOLO Geofísica Aplicada  ·  +55 (11) 99999-9999  ·  contato@scansolo.com.br"


# ── Job entry point ──────────────────────────────────────────────────────────

def handle_relatorio_job(supa: "SupabaseClient", job: dict) -> None:
    job_id: str = job["id"]
    project_id: str = job["project_id"]

    log.info("relatorio_job_start", job_id=job_id, project_id=project_id)
    supa.update_job_status(job_id, "processando")
    supa.update_project_status(project_id, "relatorio_em_andamento")

    project = supa.get_project(project_id)
    if not project:
        raise RuntimeError(f"Project {project_id} not found")

    # Targets aprovados para o relatório
    profiles = supa.get_profiles_for_project(project_id)
    run_id = supa.get_latest_run_id(project_id)
    profiles = [p for p in profiles if p.get("run_id") == run_id]

    all_targets = _get_all_targets(supa, profiles)
    relatorio_targets = _get_relatorio_targets(supa, all_targets)
    ai_map = _get_ai_map(supa, [t["id"] for t in all_targets])

    log.info("relatorio_targets", total=len(all_targets), para_relatorio=len(relatorio_targets))

    # Versão (incremental)
    existing = supa.get_report_outputs(project_id)
    version = len(existing) + 1

    # Gerar DOCX
    docx_bytes = _gen_docx(project, relatorio_targets, profiles, ai_map)

    # Upload
    docx_path = f"{project_id}/relatorio/relatorio_v{version:02d}.docx"
    supa.upload_file(STORAGE_BUCKET, docx_path, docx_bytes,
                     "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    log.info("relatorio_docx_uploaded", path=docx_path)

    # Inserir record
    supa.insert_report_output({
        "project_id": project_id,
        "version": version,
        "docx_dropbox_path": docx_path,
        "docx_storage_url": f"{SUPABASE_URL}/storage/v1/object/public/{STORAGE_BUCKET}/{docx_path}",
        "status": "gerado",
        "dados_usados_json": {
            "n_targets": len(relatorio_targets),
            "n_profiles": len(profiles),
            "run_id": run_id,
        },
    })

    supa.update_job_status(job_id, "concluido")
    supa.update_project_status(project_id, "relatorio_gerado")
    log.info("relatorio_job_done", job_id=job_id, version=version, targets=len(relatorio_targets))


# ── Data helpers ─────────────────────────────────────────────────────────────

def _get_all_targets(supa: "SupabaseClient", profiles: list[dict]) -> list[dict]:
    if not profiles:
        return []
    profile_ids = [p["id"] for p in profiles]
    r = supa._client.table("detected_targets").select("*").in_("profile_id", profile_ids).order("rank").execute()
    return r.data or []


def _get_relatorio_targets(supa: "SupabaseClient", targets: list[dict]) -> list[dict]:
    """Returns targets where vai_para_relatorio=True from technical_reviews."""
    if not targets:
        return []
    target_ids = [t["id"] for t in targets]
    r = supa._client.table("technical_reviews").select("*").in_("target_id", target_ids).execute()
    reviews = {rv["target_id"]: rv for rv in (r.data or [])}

    result = []
    for t in targets:
        rv = reviews.get(t["id"], {})
        if rv.get("vai_para_relatorio"):
            result.append({**t, "review": rv, "tipo_final": rv.get("tipo_final") or t.get("tipo_material")})
    return result


def _get_ai_map(supa: "SupabaseClient", target_ids: list[str]) -> dict[str, dict]:
    if not target_ids:
        return {}
    r = supa._client.table("ai_interpretations").select(
        "target_id, ia_tipo_sugerido, ia_descricao, ia_justificativa_tecnica, ia_confianca"
    ).in_("target_id", target_ids).execute()
    return {a["target_id"]: a for a in (r.data or [])}


# ── DOCX generation ──────────────────────────────────────────────────────────

def _gen_docx(
    project: dict[str, Any],
    targets: list[dict[str, Any]],
    profiles: list[dict[str, Any]],
    ai_map: dict[str, dict],
) -> bytes:
    doc = Document()

    # Page setup A4
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.0)
    sec.top_margin = Cm(3.0)
    sec.bottom_margin = Cm(2.5)

    _setup_styles(doc)
    _add_cover(doc, project)
    _add_header_footer(doc, project)
    doc.add_page_break()

    _add_toc(doc)
    doc.add_page_break()

    _add_h1(doc, "1. APRESENTAÇÃO")
    _add_body(doc, _text_apresentacao(project))

    _add_h1(doc, "2. OBJETIVO")
    _add_body(doc, _text_objetivo(project))

    _add_h1(doc, "3. METODOLOGIA")
    _add_h2(doc, "3.1 Georadar (GPR)")
    _add_body(doc, _BOILERPLATE_GPR)
    _add_h2(doc, "3.2 Pipe Locator")
    _add_body(doc, _BOILERPLATE_PIPE_LOCATOR if project.get("tem_pipe_locator") else
              "Pipe Locator não utilizado neste levantamento.")

    _add_h1(doc, "4. LEVANTAMENTO DE CAMPO")
    _add_h2(doc, "4.1 Área Levantada")
    _add_body(doc, _text_area(project))
    _add_h2(doc, "4.2 Aspectos Técnicos")
    _add_body(doc, _text_aspectos(project))

    _add_h1(doc, "5. RESULTADOS")
    _add_body(doc, _BOILERPLATE_DISCLAIMER)
    _add_resultados(doc, targets, profiles, ai_map)

    _add_h1(doc, "6. CONCLUSÃO")
    _add_body(doc, _BOILERPLATE_CONCLUSAO)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Seções variáveis ─────────────────────────────────────────────────────────

def _text_apresentacao(p: dict) -> str:
    cliente = p.get("cliente") or "Cliente"
    contato = p.get("contato_nome") or ""
    codigo = p.get("codigo_projeto") or p.get("codigo_interno") or p.get("nome")
    nome = p.get("nome") or ""
    local = p.get("local") or ""
    estado = p.get("estado") or ""

    dest = f"À {cliente}"
    if contato:
        dest += f" / A/C: {contato}"

    return (
        f"{dest}\n\n"
        "O presente relatório contém o conjunto de análises realizadas após o levantamento "
        "com Georadar (GPR) e Pipe Locator para mapeamento de interferências subterrâneas, "
        f"em trechos do PROJETO {codigo} – {nome} na {cliente}, localizada em "
        f"{local} / {estado}."
    )


def _text_objetivo(p: dict) -> str:
    local = p.get("local") or "área do projeto"
    return (
        f"O objetivo do levantamento foi identificar, mapear e registrar possíveis "
        f"interferências subterrâneas em {local}, fornecendo subsídios técnicos para "
        "planejamento de obras, escavações ou manutenção de infraestrutura subterrânea."
    )


def _text_area(p: dict) -> str:
    codigo = p.get("codigo_projeto") or p.get("codigo_interno") or p.get("nome")
    nome = p.get("nome") or ""
    cliente = p.get("cliente") or ""
    local = p.get("local") or ""
    estado = p.get("estado") or ""
    area = p.get("area_m2")
    area_str = f"{area:.0f}m²" if area else "[área a confirmar]m²"

    return (
        f"O levantamento foi realizado em trecho de {area_str} referente ao PROJETO "
        f"{codigo} – {nome} na {cliente}, localizada em {local} / {estado}. "
        "Foram executadas linhas de sondagem Georadar e Pipe Locator em forma de malha "
        "cobrindo a área solicitada, conforme norma ABNT 15935."
    )


def _text_aspectos(p: dict) -> str:
    mhz = p.get("antena_freq_mhz") or 270
    # Depth capacity is roughly inversely proportional to frequency
    depth_cap = "3,0" if mhz <= 270 else "2,0" if mhz <= 400 else "1,5"
    return (
        f"Foi utilizada antena GSSI de {mhz}MHz, com capacidade de investigação até "
        f"aproximadamente {depth_cap}m de profundidade em solos argilosos e maior em solos "
        "secos e arenosos. A aquisição dos dados foi realizada com parâmetros ajustados "
        "para as condições locais do terreno, garantindo máxima resolução e profundidade "
        "de investigação adequadas ao objetivo do levantamento."
    )


def _add_resultados(
    doc: Document,
    targets: list[dict],
    profiles: list[dict],
    ai_map: dict[str, dict],
) -> None:
    if not targets:
        _add_body(doc, "Nenhuma interferência identificada para inclusão no relatório.")
        return

    profile_map = {p["id"]: p for p in profiles}

    # Group targets by profile for image insertion
    by_profile: dict[str, list[dict]] = {}
    for t in targets:
        pid = t.get("profile_id", "unknown")
        by_profile.setdefault(pid, []).append(t)

    for profile_id, ptargets in by_profile.items():
        prof = profile_map.get(profile_id, {})
        arquivo = prof.get("arquivo_dzt", "")

        if arquivo:
            p = doc.add_paragraph()
            p.add_run(f"Perfil: {arquivo}").bold = True
            p.style = doc.styles["Normal"]

        for t in ptargets:
            ai = ai_map.get(t["id"], {})
            rv = t.get("review", {})
            tipo_final = t.get("tipo_final") or ai.get("ia_tipo_sugerido") or "desconhecido"
            tipo_label = _TIPO_LABEL.get(tipo_final, tipo_final.replace("_", " ").title())
            depth_m = t.get("depth_m") or 0
            diam_m = t.get("diam_est_m") or 0
            diam_mm = int(round(diam_m * 1000))
            depth_br = f"{depth_m:.2f}".replace(".", ",")
            rank = t.get("rank", "?")

            # Header line
            ph = doc.add_paragraph(style="Normal")
            run = ph.add_run(f"Interferência #{rank}: {tipo_label}")
            run.bold = True
            run.font.size = Pt(11)

            # Specs line
            ps = doc.add_paragraph(style="Normal")
            ps.add_run(
                f"DIÂMETRO APARENTE = {diam_mm}mm⌀  |  PROF. GS = {depth_br}m"
            )
            ps.paragraph_format.left_indent = Cm(0.5)

            # IA description
            descricao = ai.get("ia_descricao") or rv.get("observacao")
            if descricao:
                pd = doc.add_paragraph(style="Normal")
                pd.add_run(descricao).italic = True
                pd.paragraph_format.left_indent = Cm(0.5)

            doc.add_paragraph()  # spacing

        # Radargram image
        img_url = prof.get("imagem_anotada_url") or prof.get("imagem_processada_url")
        if img_url:
            img_bytes = _download_image(img_url)
            if img_bytes:
                try:
                    p_img = doc.add_paragraph(style="Normal")
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p_img.add_run()
                    run.add_picture(io.BytesIO(img_bytes), width=Cm(15))
                    cap = doc.add_paragraph(style="Normal")
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap.add_run(f"Radargrama: {arquivo}").italic = True
                    cap.runs[0].font.size = Pt(9)
                except Exception as exc:
                    log.warning("relatorio_image_failed", arquivo=arquivo, error=str(exc))

        doc.add_page_break()


# ── DOCX formatting helpers ───────────────────────────────────────────────────

def _setup_styles(doc: Document) -> None:
    """Ensure basic styles exist with sensible fonts."""
    for style_name in ("Normal", "Heading 1", "Heading 2"):
        try:
            s = doc.styles[style_name]
            s.font.name = "Arial"
        except Exception:
            pass


def _add_cover(doc: Document, project: dict) -> None:
    TITLE = (
        "Relatório de Levantamento Geofísico com Georadar (GPR) e Pipe Locator, "
        "para Mapeamento de Estruturas e Interferências Subterrâneas"
    )
    codigo = project.get("codigo_projeto") or project.get("codigo_interno") or project.get("nome", "")
    local = project.get("local") or ""
    estado = project.get("estado") or ""
    data_str = date.today().strftime("%B de %Y").capitalize()

    for _ in range(6):
        doc.add_paragraph()

    p_title = doc.add_paragraph(TITLE)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.runs[0].bold = True
    p_title.runs[0].font.size = Pt(16)

    doc.add_paragraph()

    for text, size in [
        (f"Código: {codigo}", 13),
        (f"Local: {local} / {estado}", 12),
        (f"Data: {data_str}", 12),
    ]:
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(size)

    for _ in range(8):
        doc.add_paragraph()

    p_footer = doc.add_paragraph(_RODAPE_FIXO)
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_footer.runs[0].font.size = Pt(9)
    p_footer.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def _add_header_footer(doc: Document, project: dict) -> None:
    codigo = project.get("codigo_projeto") or project.get("codigo_interno") or project.get("nome", "")
    local = project.get("local") or ""
    estado = project.get("estado") or ""
    dwg_ref = f"{codigo}-PLT"

    sec = doc.sections[0]

    # Header
    hdr = sec.header
    hdr.is_linked_to_previous = False
    hdr_p = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
    hdr_p.clear()
    hdr_p.text = f"Relatório de Levantamento  |  Desenhos: {dwg_ref}  |  Revisão: 00"
    hdr_p.runs[0].font.size = Pt(8)
    hdr_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    hdr_p2 = hdr.add_paragraph()
    hdr_p2.text = (
        f"Levantamento Geofísico com Georadar (GPR) e Pipe Locator, "
        f"Mapeamento de Estruturas e Interferências Subterrâneas, em {local}, {estado}"
    )
    hdr_p2.runs[0].font.size = Pt(8)
    hdr_p2.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Footer with page number
    ftr = sec.footer
    ftr.is_linked_to_previous = False
    ftr_p = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
    ftr_p.clear()
    ftr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ftr_p.add_run("Página ")
    run.font.size = Pt(9)
    _add_field(ftr_p, "PAGE")
    ftr_p.add_run(" / ").font.size = Pt(9)
    _add_field(ftr_p, "NUMPAGES")


def _add_field(paragraph: Any, field_name: str) -> None:
    run = paragraph.add_run()
    run.font.size = Pt(9)
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar1)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field_name} "
    run._r.append(instr)
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar2)


def _add_toc(doc: Document) -> None:
    _add_h1(doc, "SUMÁRIO")
    toc_items = [
        ("1.", "Apresentação"),
        ("2.", "Objetivo"),
        ("3.", "Metodologia"),
        ("3.1", "Georadar (GPR)"),
        ("3.2", "Pipe Locator"),
        ("4.", "Levantamento de Campo"),
        ("4.1", "Área Levantada"),
        ("4.2", "Aspectos Técnicos"),
        ("5.", "Resultados"),
        ("6.", "Conclusão"),
    ]
    for num, title in toc_items:
        p = doc.add_paragraph(style="Normal")
        p.add_run(f"{num}  {title}")
        indent = Cm(0.5) if "." in num and num != num.split(".")[0] + "." else Cm(0)
        p.paragraph_format.left_indent = indent


def _add_h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="Heading 1")
    p.runs[0].font.size = Pt(13)
    p.runs[0].bold = True


def _add_h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="Heading 2")
    p.runs[0].font.size = Pt(11)
    p.runs[0].bold = True


def _add_body(doc: Document, text: str) -> None:
    for para in text.split("\n\n"):
        if para.strip():
            p = doc.add_paragraph(para.strip(), style="Normal")
            p.runs[0].font.size = Pt(11)
            p.paragraph_format.space_after = Pt(6)


# ── Image download ────────────────────────────────────────────────────────────

def _download_image(url: str) -> bytes | None:
    if not url or not SUPABASE_URL:
        return None
    for bucket in ("gpr-images", "gpr-tabelas"):
        prefix = f"{SUPABASE_URL}/storage/v1/object/public/{bucket}/"
        if url.startswith(prefix):
            path = url[len(prefix):]
            try:
                import truststore
                truststore.inject_into_ssl()
                from clients.supabase_client import SupabaseClient
                supa = SupabaseClient()
                return supa.download_file(bucket, path)
            except Exception as exc:
                log.warning("relatorio_img_download_failed", path=path, error=str(exc))
                return None
    return None
